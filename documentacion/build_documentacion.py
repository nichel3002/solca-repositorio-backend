from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


BASE = Path(__file__).resolve().parent
OUT = BASE / "documentacion_avance_solca_repositorio_clinico.docx"
ASSETS = BASE / "assets"

BLUE = RGBColor(31, 77, 120)
MID_BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(24, 36, 48)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
INK = "#17324D"
ACCENT = "#2E74B5"
GREEN = "#2F6F3E"
SOFT_BLUE = "#E8EEF5"
SOFT_GREEN = "#E9F4EC"
SOFT_GRAY = "#F5F7FA"
RISK = "#B42318"


def font(size, bold=False):
    names = ["arialbd.ttf" if bold else "arial.ttf", "calibrib.ttf" if bold else "calibri.ttf"]
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw, text, font_obj, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        box = draw.textbbox((0, 0), candidate, font=font_obj)
        if box[2] - box[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_text_center(draw, xy, text, font_obj, fill=INK, max_width=None, line_gap=6):
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    if max_width:
        split_lines = []
        for line in lines:
            split_lines.extend(wrap_text(draw, line, font_obj, max_width))
        lines = split_lines
    heights = []
    widths = []
    for line in lines:
        box = draw.textbbox((0, 0), line, font=font_obj)
        widths.append(box[2] - box[0])
        heights.append(box[3] - box[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font_obj, fill=fill)
        y += h + line_gap


def rounded(draw, xy, fill, outline=ACCENT, width=3, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=ACCENT, width=4):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 18, ey - 9), (ex - direction * 18, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - direction * 18), (ex + 9, ey - direction * 18)]
    draw.polygon(points, fill=color)


def save_diagram(name, draw_func, size=(1800, 1050)):
    ASSETS.mkdir(exist_ok=True)
    path = ASSETS / name
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw_func(draw, size)
    img.save(path, quality=95)
    return path


def build_diagram_architecture():
    def render(draw, size):
        title = font(42, True)
        h = font(28, True)
        body = font(22)
        small = font(18)
        draw_text_center(draw, (80, 30, 1720, 90), "Arquitectura cloud preliminar - Repositorio Clinico Regional SOLCA", title, ACCENT)

        rounded(draw, (60, 130, 1740, 940), "#FFFFFF", "#C8D6E5", 4, 28)
        draw_text_center(draw, (80, 145, 1720, 190), "Nube / plataforma PaaS - Docker / Railway", h, INK)

        # Sites
        sites = [("SOLCA Cuenca", 110), ("SOLCA Manabi", 310), ("SOLCA Guayaquil", 510)]
        for label, y in sites:
            rounded(draw, (110, y, 340, y + 110), SOFT_GREEN, GREEN, 3, 16)
            draw_text_center(draw, (125, y + 15, 325, y + 95), label + "\nSistema local", body, GREEN, 180)
            arrow(draw, (345, y + 55), (450, 405), GREEN, 4)

        rounded(draw, (450, 335, 720, 475), SOFT_BLUE, ACCENT, 4, 18)
        draw_text_center(draw, (470, 350, 700, 460), "Frontend Angular\nLogin JWT\nConsulta regional", body, INK, 220)
        arrow(draw, (720, 405), (840, 405), ACCENT, 5)

        rounded(draw, (840, 335, 1110, 475), "#FFF7E6", "#C47F00", 4, 18)
        draw_text_center(draw, (860, 350, 1090, 460), "API REST\nAuthorization:\nBearer Token", body, "#7A4C00", 220)
        arrow(draw, (1110, 405), (1230, 405), ACCENT, 5)

        rounded(draw, (1230, 315, 1580, 500), SOFT_BLUE, ACCENT, 4, 18)
        draw_text_center(draw, (1250, 330, 1560, 485), "Repositorio Clinico Regional\nGET /repositorio/paciente/{id}\nAgrega datos sin acceder a BD externas", body, INK, 300)

        services = [
            ("Paciente Maestro", "paciente_maestro_db", 850, 610),
            ("Consulta Clinica", "consulta_clinica_db", 1100, 610),
            ("Laboratorio", "laboratorio_clinico_db", 1350, 610),
            ("Imagenologia/PACS", "imagenologia_db", 850, 790),
        ]
        for label, db, x, y in services:
            rounded(draw, (x, y, x + 220, y + 120), SOFT_GRAY, "#73808C", 3, 14)
            draw_text_center(draw, (x + 10, y + 10, x + 210, y + 72), label, small, INK, 185)
            draw.rounded_rectangle((x + 35, y + 78, x + 185, y + 110), radius=10, fill="#FFFFFF", outline="#73808C", width=2)
            draw_text_center(draw, (x + 40, y + 80, x + 180, y + 108), db, font(14), "#46515C", 135)
            arrow(draw, (1405, 500), (x + 110, y), ACCENT, 3)

        rounded(draw, (1100, 790, 1320, 910), "#F0F7FF", "#4A90E2", 3, 14)
        draw_text_center(draw, (1115, 805, 1305, 895), "Storage cloud\nimagenes, informes,\narchivos clinicos", small, INK, 180)
        arrow(draw, (1405, 500), (1210, 790), "#4A90E2", 3)

        draw_text_center(draw, (80, 965, 1720, 1020), "Principio clave: cada microservicio conserva su propia base de datos; la integracion se hace por APIs REST.", font(22, True), GREEN)

    return save_diagram("diagrama_arquitectura_cloud_solca.png", render)


def build_diagram_integration():
    def render(draw, size):
        title = font(42, True)
        h = font(24, True)
        body = font(20)
        draw_text_center(draw, (80, 30, 1720, 90), "Integracion REST del Repositorio Clinico Regional", title, ACCENT)
        actors = [
            ("Usuario / Postman", 90),
            ("Frontend Angular", 360),
            ("Repositorio\nRegional", 650),
            ("Paciente\nMaestro", 980),
            ("Consulta\nClinica", 1210),
            ("Laboratorio", 1440),
            ("Imagenologia", 1440),
        ]
        x_positions = [120, 390, 690, 990, 1210, 1440, 1440]
        y_boxes = [160, 160, 160, 160, 160, 160, 600]
        for (label, _), x, y in zip(actors, x_positions, y_boxes):
            rounded(draw, (x, y, x + 190, y + 90), SOFT_BLUE if x != 1440 or y != 600 else SOFT_GREEN, ACCENT if y != 600 else GREEN, 3, 14)
            draw_text_center(draw, (x + 10, y + 12, x + 180, y + 78), label, h, INK, 170)
            draw.line([(x + 95, y + 90), (x + 95, 910)], fill="#CED8E5", width=3)

        steps = [
            ((310, 300), (390, 300), "1. Login /auth/login"),
            ((580, 365), (690, 365), "2. GET /repositorio/paciente/REG-0001"),
            ((880, 430), (990, 430), "3. GET /pacientes/{id}"),
            ((880, 500), (1210, 500), "4. GET /consultas/paciente/{id}"),
            ((880, 570), (1440, 570), "5. GET /laboratorio/paciente/{id}"),
            ((880, 640), (1440, 640), "6. GET /imagenes/paciente/{id}"),
        ]
        for start, end, label in steps:
            arrow(draw, start, end, ACCENT, 4)
            draw.text((start[0] + 8, start[1] - 30), label, font=body, fill=INK)

        returns = [
            ((990, 465), (880, 465), "paciente {}"),
            ((1210, 535), (880, 535), "consultas []"),
            ((1440, 605), (880, 605), "laboratorio []"),
            ((1440, 675), (880, 675), "imagenes []"),
            ((690, 770), (580, 770), "JSON consolidado"),
            ((390, 835), (310, 835), "Vista regional del paciente"),
        ]
        for start, end, label in returns:
            arrow(draw, start, end, GREEN, 3)
            draw.text((end[0] + 10, start[1] - 28), label, font=body, fill=GREEN)

        rounded(draw, (760, 735, 1120, 885), "#FFF7E6", "#C47F00", 3, 16)
        draw_text_center(draw, (780, 750, 1100, 870), "Manejo de errores:\nsi un microservicio no responde,\nla seccion vuelve vacia y se agrega\nen 'errores'.", body, "#7A4C00", 310)

    return save_diagram("diagrama_integracion_rest_solca.png", render)


def build_diagram_data_model():
    def render(draw, size):
        title = font(42, True)
        h = font(25, True)
        body = font(19)
        draw_text_center(draw, (80, 30, 1720, 90), "Modelo de datos: Paciente Maestro y bases independientes", title, ACCENT)

        central = (640, 150, 1160, 390)
        rounded(draw, central, SOFT_BLUE, ACCENT, 4, 18)
        draw_text_center(draw, (660, 165, 1140, 205), "Paciente Maestro Regional", h, INK)
        fields = [
            "idPacienteRegional (PK)",
            "cedula",
            "nombres, apellidos",
            "fechaNacimiento, sexo",
            "sedeOrigen",
            "telefono, direccion",
        ]
        y = 220
        for item in fields:
            draw.text((700, y), item, font=body, fill=INK)
            y += 27

        entities = [
            ("Consulta Clinica", ["id", "idPacienteRegional", "fechaConsulta", "diagnostico", "tratamiento"], (150, 520, 500, 800), "consulta_clinica_db"),
            ("Laboratorio Clinico", ["id", "idPacienteRegional", "tipoExamen", "resultado", "fechaResultado"], (540, 520, 890, 800), "laboratorio_clinico_db"),
            ("Imagenologia / PACS", ["id", "idPacienteRegional", "tipoEstudio", "urlArchivo", "informe"], (930, 520, 1280, 800), "imagenologia_db"),
            ("Registro Repositorio", ["id", "idPacienteRegional", "fechaConsulta", "estado", "mensajeError"], (1320, 520, 1670, 800), "repositorio_clinico_regional_db"),
        ]
        for name, attrs, xy, db in entities:
            rounded(draw, xy, SOFT_GRAY, "#73808C", 3, 14)
            draw_text_center(draw, (xy[0] + 10, xy[1] + 15, xy[2] - 10, xy[1] + 55), name, h, INK, xy[2] - xy[0] - 30)
            y = xy[1] + 75
            for attr in attrs:
                draw.text((xy[0] + 35, y), attr, font=body, fill=INK)
                y += 30
            draw.rounded_rectangle((xy[0] + 30, xy[3] - 55, xy[2] - 30, xy[3] - 18), radius=10, fill="#FFFFFF", outline="#73808C", width=2)
            draw_text_center(draw, (xy[0] + 35, xy[3] - 52, xy[2] - 35, xy[3] - 20), db, font(15), "#46515C", xy[2] - xy[0] - 80)
            arrow(draw, ((central[0] + central[2]) // 2, central[3]), ((xy[0] + xy[2]) // 2, xy[1]), ACCENT, 3)

        rounded(draw, (320, 865, 1480, 970), SOFT_GREEN, GREEN, 3, 16)
        draw_text_center(draw, (340, 880, 1460, 955), "Relacion logica por idPacienteRegional. No existe acceso directo entre bases; la union ocurre en el Repositorio mediante llamadas REST.", font(22, True), GREEN, 1080)

    return save_diagram("diagrama_modelo_datos_solca.png", render)


def build_diagram_components():
    def render(draw, size):
        title = font(42, True)
        h = font(25, True)
        body = font(20)
        draw_text_center(draw, (80, 30, 1720, 90), "Componentes cloud identificados para el proyecto", title, ACCENT)
        comps = [
            ("IaaS", "Contenedores y recursos de computo para ejecutar servicios si se requiere control de infraestructura.", (120, 170, 520, 390), "#F0F7FF", ACCENT),
            ("PaaS", "Railway para desplegar microservicios Spring Boot y frontend con variables de entorno.", (700, 170, 1100, 390), "#EAF7EF", GREEN),
            ("DBaaS", "PostgreSQL administrado con una base por microservicio y backups.", (1280, 170, 1680, 390), "#FFF7E6", "#C47F00"),
            ("SaaS", "GitHub, Docker Hub y Postman para codigo, imagenes y pruebas.", (330, 560, 730, 780), "#F5F7FA", "#73808C"),
            ("Storage cloud", "Repositorio de archivos no estructurados: imagenes PACS, informes y documentos clinicos.", (1070, 560, 1470, 780), "#FCEEEE", RISK),
        ]
        for label, detail, xy, fill, outline in comps:
            rounded(draw, xy, fill, outline, 4, 18)
            draw_text_center(draw, (xy[0] + 20, xy[1] + 22, xy[2] - 20, xy[1] + 70), label, h, outline)
            draw_text_center(draw, (xy[0] + 30, xy[1] + 85, xy[2] - 30, xy[3] - 30), detail, body, INK, xy[2] - xy[0] - 70)

        center = (760, 420, 1040, 520)
        rounded(draw, center, SOFT_BLUE, ACCENT, 4, 18)
        draw_text_center(draw, center, "Repositorio Clinico\nRegional SOLCA", h, INK, 250)
        for _, _, xy, _, outline in comps:
            arrow(draw, ((xy[0] + xy[2]) // 2, xy[3] if xy[1] < center[1] else xy[1]), ((center[0] + center[2]) // 2, center[1] if xy[1] < center[1] else center[3]), outline, 3)

    return save_diagram("diagrama_componentes_cloud_solca.png", render)


def generate_diagrams():
    return [
        ("Figura 1. Arquitectura cloud preliminar del Repositorio Clinico Regional SOLCA.", build_diagram_architecture()),
        ("Figura 2. Integracion REST para el endpoint consolidado del repositorio.", build_diagram_integration()),
        ("Figura 3. Modelo de datos del Paciente Maestro Regional y bases independientes.", build_diagram_data_model()),
        ("Figura 4. Componentes cloud IaaS, PaaS, SaaS, DBaaS y storage cloud.", build_diagram_components()),
    ]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2F3"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True, color=BLUE)
        if widths:
            set_cell_width(cell, widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                set_cell_width(cells[idx], widths[idx])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, "B7C9E2")
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = BLUE
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = BLUE
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Proyecto Reto SOLCA - Repositorio Clinico Regional")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 100, 110)


def add_cover(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Documentacion del Proyecto Reto").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Repositorio Clinico Regional SOLCA")
    r.font.size = Pt(16)
    r.font.color.rgb = BLUE
    r.bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Avance 1 - Diseno inicial y microservicios base\n").bold = True
    meta.add_run("Avance 2 - Repositorio Clinico Regional e integracion\n")
    meta.add_run("Sedes consideradas: SOLCA Cuenca, SOLCA Manabi y SOLCA Guayaquil")
    doc.add_paragraph()
    add_callout(
        doc,
        "Resumen ejecutivo",
        "El proyecto plantea una arquitectura basada en microservicios Spring Boot y frontend Angular no standalone para integrar informacion clinica regional. Cada microservicio mantiene su propia base PostgreSQL y el Repositorio Clinico Regional consolida datos mediante APIs REST protegidas con JWT, sin acceso directo entre bases de datos."
    )
    add_table(
        doc,
        ["Elemento", "Estado"],
        [
            ["Backend", "Cinco microservicios Spring Boot independientes"],
            ["Frontend", "Angular modular con login JWT"],
            ["Base de datos", "PostgreSQL con bases independientes por dominio"],
            ["Integracion", "REST entre servicios y respuesta JSON consolidada"],
            ["Evidencia", "GitHub, Docker, Postman, pgAdmin y capturas sugeridas"],
        ],
        [2200, 7160],
    )
    doc.add_page_break()


def add_design_sections(doc):
    doc.add_heading("1. Analisis del problema actual", level=1)
    doc.add_paragraph(
        "SOLCA Cuenca, Manabi y Guayaquil atienden pacientes oncologicos que pueden requerir continuidad clinica entre sedes. El problema principal es la fragmentacion de informacion: datos demograficos, consultas, resultados de laboratorio e imagenologia pueden quedar distribuidos en sistemas locales, dificultando una vision regional del paciente."
    )
    add_bullets(doc, [
        "Riesgo de duplicidad de registros de paciente cuando cada sede administra su propia informacion.",
        "Demoras en la consulta de antecedentes clinicos, resultados de laboratorio e informes de imagenologia.",
        "Necesidad de integrar informacion estructurada y no estructurada sin mezclar directamente las bases de datos.",
        "Requerimiento de trazabilidad, seguridad y control de acceso para informacion clinica sensible.",
        "Necesidad de una arquitectura escalable que pueda operar en nube y permitir despliegues independientes.",
    ])

    doc.add_heading("2. Diagrama preliminar de arquitectura cloud", level=1)
    add_callout(
        doc,
        "Vista conceptual",
        "La arquitectura se organiza en una aplicacion frontend Angular, cinco microservicios Spring Boot, bases PostgreSQL independientes y un repositorio regional que actua como agregador REST. Cada sede conserva su sistema local, mientras el repositorio regional consulta los dominios clinicos mediante APIs."
    )
    add_table(
        doc,
        ["Capa", "Componentes", "Funcion"],
        [
            ["Presentacion", "Angular no standalone", "Permite login JWT y consulta de historia regional."],
            ["Servicios", "Paciente, Consulta, Laboratorio, Imagenologia, Repositorio", "Expone APIs REST por dominio clinico."],
            ["Datos", "PostgreSQL por microservicio", "Mantiene independencia de esquemas y evita acoplamiento directo."],
            ["Seguridad", "JWT", "Protege endpoints y permite autorizacion por token."],
            ["Operacion", "Docker / Railway", "Facilita despliegue por servicio y escalabilidad."],
        ],
        [1500, 3100, 4760],
    )

    doc.add_heading("3. Modelo del Paciente Maestro Regional", level=1)
    doc.add_paragraph(
        "El Paciente Maestro Regional unifica la identidad minima del paciente para que las sedes puedan relacionar consultas, examenes e imagenes con un mismo identificador regional."
    )
    add_table(
        doc,
        ["Campo", "Descripcion"],
        [
            ["idPacienteRegional", "Identificador regional unico, por ejemplo REG-0001."],
            ["cedula", "Documento de identidad del paciente."],
            ["nombres / apellidos", "Datos personales principales."],
            ["sedeOrigen", "SOLCA Cuenca, SOLCA Manabi o SOLCA Quito."],
            ["fechaNacimiento / sexo", "Datos demograficos."],
            ["direccion / telefono", "Datos de contacto."],
        ],
        [2500, 6860],
    )

    doc.add_heading("4. Bases de datos independientes", level=1)
    add_table(
        doc,
        ["Microservicio", "Base de datos", "Tabla principal"],
        [
            ["Paciente Maestro", "paciente_maestro_db", "pacientes"],
            ["Consulta Clinica", "consulta_clinica_db", "consultas_clinicas"],
            ["Laboratorio Clinico", "laboratorio_clinico_db", "resultados_laboratorio"],
            ["Imagenologia / PACS", "imagenologia_db", "estudios_imagenologia"],
            ["Repositorio Regional", "repositorio_clinico_regional_db", "registro_consultas_repositorio"],
        ],
        [2700, 3300, 3360],
    )

    doc.add_heading("5. Componentes cloud identificados", level=1)
    add_table(
        doc,
        ["Categoria", "Aplicacion en el proyecto"],
        [
            ["IaaS", "Maquinas virtuales o contenedores para ejecutar servicios backend y base de datos si se requiere control de infraestructura."],
            ["PaaS", "Railway o plataforma similar para desplegar microservicios sin administrar servidores directamente."],
            ["SaaS", "GitHub para repositorios, Docker Hub para imagenes y Postman para pruebas documentadas."],
            ["DBaaS", "PostgreSQL administrado por la plataforma cloud para cada base independiente."],
            ["Storage cloud", "Almacenamiento de imagenes DICOM/PACS o documentos clinicos no estructurados."],
        ],
        [1800, 7560],
    )


def add_implementation_sections(doc):
    doc.add_heading("6. Implementacion del Avance 1", level=1)
    add_table(
        doc,
        ["Microservicio", "Puerto", "Endpoints minimos"],
        [
            ["Paciente Maestro Regional", "8081", "GET /pacientes, GET /pacientes/{id}, POST /pacientes"],
            ["Consulta Clinica", "8082", "GET /consultas, GET /consultas/paciente/{id}, POST /consultas"],
            ["Laboratorio Clinico", "8083", "GET /laboratorio, GET /laboratorio/paciente/{id}, POST /laboratorio"],
            ["Imagenologia / PACS", "8084", "GET /imagenes, GET /imagenes/paciente/{id}, POST /imagenes"],
            ["Repositorio Clinico Regional", "8085", "GET /repositorio/paciente/{id}, GET /repositorio/registros"],
        ],
        [2700, 900, 5760],
    )
    add_callout(
        doc,
        "Autenticacion",
        "Todos los endpoints principales estan protegidos con JWT. Para las pruebas se utiliza POST /auth/login con usuario admin y clave admin123. El token se envia como Authorization: Bearer TOKEN."
    )

    doc.add_heading("7. Implementacion del Avance 2", level=1)
    doc.add_paragraph(
        "El Repositorio Clinico Regional consume los microservicios de Paciente Maestro, Consulta Clinica, Laboratorio e Imagenologia mediante REST. La respuesta consolidada cumple el contrato solicitado."
    )
    add_table(
        doc,
        ["Endpoint obligatorio", "Respuesta esperada"],
        [["GET /repositorio/paciente/{idPacienteRegional}", '{"paciente": {}, "consultas": [], "laboratorio": [], "imagenes": []}']],
        [3900, 5460],
    )
    add_bullets(doc, [
        "Las llamadas entre servicios se realizan por HTTP REST.",
        "El repositorio regional no consulta directamente las bases de datos de otros microservicios.",
        "Si un servicio no responde, el repositorio devuelve la seccion vacia y registra el error en el objeto errores.",
        "La prueba principal se realiza con REG-0001, REG-0002 y REG-0003.",
    ])

    doc.add_heading("8. Infraestructura de datos", level=1)
    add_table(
        doc,
        ["Tipo", "Uso"],
        [
            ["Datos estructurados", "Pacientes, consultas, resultados de laboratorio, metadatos de imagenes y registros de consulta."],
            ["Datos no estructurados", "Imagenes diagnosticas, informes radiologicos extendidos, documentos clinicos y archivos PACS."],
            ["DBaaS", "PostgreSQL administrado para cada microservicio."],
            ["Storage cloud", "Repositorio para imagenes o documentos clinicos asociados a PACS."],
        ],
        [2400, 6960],
    )

    doc.add_page_break()
    doc.add_heading("9. Matriz preliminar de riesgos", level=1)
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigacion"],
        [
            ["Microservicio no disponible", "Respuesta parcial del repositorio", "Timeouts, manejo de errores y monitoreo."],
            ["Duplicidad de pacientes", "Historia clinica fragmentada", "Paciente Maestro Regional y validacion por identificador."],
            ["Acceso no autorizado", "Exposicion de informacion sensible", "JWT, HTTPS en despliegue y control de roles futuro."],
            ["Falla de base de datos", "Perdida de disponibilidad del dominio", "Backups, DBaaS y replicas."],
            ["Archivos de imagen muy pesados", "Lentitud en consulta de PACS", "Storage cloud y metadatos en base estructurada."],
            ["Configuracion incorrecta entre servicios", "Errores de integracion REST", "Variables de entorno y pruebas Postman por endpoint."],
        ],
        [2600, 2600, 4160],
    )


def add_evidence(doc):
    doc.add_heading("10. Evidencia tecnica sugerida", level=1)
    add_table(
        doc,
        ["Evidencia", "Detalle"],
        [
            ["Repositorio backend", "https://github.com/nichel3002/solca-repositorio-backend.git"],
            ["Repositorio frontend", "https://github.com/nichel3002/solca-repositorio-frontend.git"],
            ["Docker Hub", "Imagenes publicadas como nichel030/solca-*-service:latest"],
            ["Postman", "Coleccion SOLCA-Repositorio-Clinico-Regional.postman_collection.json"],
            ["pgAdmin", "Host 127.0.0.1, puerto 5432, usuario solca, clave solca"],
            ["Frontend local", "http://127.0.0.1:4200"],
        ],
        [2600, 6760],
    )
    add_numbered(doc, [
        "Capturar POST /auth/login y guardar el token JWT.",
        "Capturar GET /pacientes y GET /pacientes/REG-0001.",
        "Capturar GET /consultas/paciente/REG-0001.",
        "Capturar GET /laboratorio/paciente/REG-0001.",
        "Capturar GET /imagenes/paciente/REG-0001.",
        "Capturar GET /repositorio/paciente/REG-0001 con respuesta consolidada.",
        "Capturar en pgAdmin las tablas principales de cada base independiente.",
    ])


def add_project_diagrams(doc, diagrams):
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec = doc.sections[-1]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    doc.add_heading("Anexo A. Diagramas del Repositorio Clinico Regional", level=1)
    doc.add_paragraph(
        "Los siguientes diagramas fueron elaborados a partir de la implementacion del proyecto. Representan la arquitectura cloud preliminar, la integracion REST del endpoint consolidado, el modelo de datos del Paciente Maestro Regional y los componentes cloud exigidos en los avances."
    )
    add_table(
        doc,
        ["Diagrama", "Uso en la documentacion"],
        [
            ["Arquitectura cloud", "Explica como Angular, JWT, microservicios Spring Boot, PostgreSQL y storage cloud se relacionan en una plataforma PaaS."],
            ["Integracion REST", "Muestra el flujo obligatorio GET /repositorio/paciente/{idPacienteRegional} y las llamadas a los microservicios."],
            ["Modelo de datos", "Presenta el Paciente Maestro Regional y la independencia de bases de datos por dominio clinico."],
            ["Componentes cloud", "Identifica IaaS, PaaS, SaaS, DBaaS y storage cloud como pide el enunciado."],
        ],
        [2600, 6760],
    )
    for caption, image in diagrams:
        doc.add_section(WD_SECTION.NEW_PAGE)
        sec = doc.sections[-1]
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Inches(11)
        sec.page_height = Inches(8.5)
        sec.left_margin = Inches(0.55)
        sec.right_margin = Inches(0.55)
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        doc.add_paragraph(caption).runs[0].bold = True
        if image.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Inches(9.45))
        else:
            doc.add_paragraph(f"No se encontro la imagen: {image}")


def build():
    diagrams = generate_diagrams()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_design_sections(doc)
    add_implementation_sections(doc)
    add_evidence(doc)
    add_project_diagrams(doc, diagrams)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
